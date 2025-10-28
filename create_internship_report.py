from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# Create document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Title
title = doc.add_paragraph()
title_run = title.add_run('INTERNSHIP REPORT\n')
title_run.bold = True
title_run.font.size = Pt(16)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('Quantitative Momentum Backtesting System')
subtitle_run.font.size = Pt(14)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Personal Details
details = doc.add_paragraph()
details_text = details.add_run('Student Name: [Your Name]\nInternship Duration: [Start Date] - [End Date]\nOrganization: [Company Name]\nSupervisor: [Supervisor Name]')
details_text.font.size = Pt(11)

doc.add_paragraph()

# 1. EXECUTIVE SUMMARY
heading1 = doc.add_heading('1. EXECUTIVE SUMMARY', level=1)
heading1.runs[0].font.size = Pt(14)

para1 = doc.add_paragraph(
    'This report presents the development and implementation of a comprehensive quantitative momentum backtesting system. '
    'The project involved designing and building a production-grade financial analysis tool capable of evaluating momentum-based '
    'trading strategies across a universe of 119 stocks over a 9+ year period (June 2016 - August 2025). The system successfully '
    'processes over 298,000 data points and executes thousands of simulated trades to analyze strategy performance with '
    'institutional-grade metrics and visualizations.'
)
para1.style.font.size = Pt(11)

# 2. PROJECT OBJECTIVES
heading2 = doc.add_heading('2. PROJECT OBJECTIVES', level=1)
heading2.runs[0].font.size = Pt(14)

para2 = doc.add_paragraph()
para2.style.font.size = Pt(11)
para2.add_run('The primary objectives of this project were:\n\n')

objectives = [
    'Develop a robust backtesting engine for momentum-based trading strategies',
    'Implement accurate performance metrics (CAGR, Sharpe Ratio, Maximum Drawdown, Hit Ratio)',
    'Create an interactive web-based user interface for strategy analysis',
    'Enable comparative analysis across different price types (adjusted vs. raw closing prices)',
    'Provide detailed quarterly performance tracking and portfolio selection insights',
    'Ensure scalability and accuracy for institutional-grade financial analysis'
]

for obj in objectives:
    para2.add_run(f'• {obj}\n')

# 3. METHODOLOGY & TECHNICAL IMPLEMENTATION
heading3 = doc.add_heading('3. METHODOLOGY & TECHNICAL IMPLEMENTATION', level=1)
heading3.runs[0].font.size = Pt(14)

# 3.1 System Architecture
subheading1 = doc.add_heading('3.1 System Architecture', level=2)
subheading1.runs[0].font.size = Pt(12)

para3 = doc.add_paragraph(
    'The system was built using a modular architecture with clear separation of concerns:'
)
para3.style.font.size = Pt(11)

para3_1 = doc.add_paragraph()
para3_1.style.font.size = Pt(11)
components = [
    'Core Engine (src/core/engine.py): Executes backtest simulations with daily mark-to-market calculations',
    'Data Loader (src/core/data_loader.py): Handles CSV parsing, date format detection, and price matrix construction',
    'Analytics Module (src/core/analytics.py): Computes CAGR, Sharpe ratio, drawdown, and quarterly metrics',
    'Strategy Module (src/core/strategy.py): Implements momentum ranking and portfolio selection logic',
    'Visualization Layer (src/core/plotting.py): Creates interactive charts using Plotly',
    'Web Interface (src/app/streamlit_app.py): Streamlit-based UI for user interaction'
]

for comp in components:
    para3_1.add_run(f'• {comp}\n')

# 3.2 Trading Strategy Implementation
subheading2 = doc.add_heading('3.2 Trading Strategy Implementation', level=2)
subheading2.runs[0].font.size = Pt(12)

para4 = doc.add_paragraph()
para4.style.font.size = Pt(11)
para4.add_run('The momentum strategy operates as follows:\n\n')

strategy_points = [
    'Universe: 119 stocks with daily price data from October 2015 to August 2025',
    'Lookback Period: 3 months for momentum calculation',
    'Selection Criteria: Top 24 stocks ranked by highest 3-month returns',
    'Rebalancing Frequency: Quarterly (March 31, June 30, September 30, December 31)',
    'Portfolio Allocation: Equal-weight allocation across selected stocks',
    'Cost Model: Transaction costs (10 bps) and market impact slippage (5 bps)',
    'Starting Capital: $1,000,000'
]

for point in strategy_points:
    para4.add_run(f'• {point}\n')

# 3.3 Technical Challenges & Solutions
subheading3 = doc.add_heading('3.3 Technical Challenges & Solutions', level=2)
subheading3.runs[0].font.size = Pt(12)

para5 = doc.add_paragraph()
para5.style.font.size = Pt(11)

challenges = [
    ('Date Format Handling', 'Implemented flexible date parsing supporting multiple formats (DD-MM-YYYY, YYYY-MM-DD, Excel serial dates)'),
    ('Position Holding Bug', 'Fixed critical issue where portfolio positions were incorrectly liquidated between rebalances'),
    ('Cash Accounting', 'Corrected double-counting of equity by properly tracking cash vs. invested capital'),
    ('Quarterly Returns Calculation', 'Redesigned to calculate from first-to-last business day of each quarter rather than compounding daily returns'),
    ('Cache Management', 'Implemented session state clearing to prevent stale results when parameters change'),
    ('Performance Optimization', 'Optimized data structures for handling 2,500+ days × 119 tickers efficiently')
]

for challenge, solution in challenges:
    para5.add_run(f'• ').bold = True
    para5.add_run(f'{challenge}: ')
    para5.add_run(f'{solution}\n')

# Page break
doc.add_page_break()

# 4. RESULTS & KEY FINDINGS
heading4 = doc.add_heading('4. RESULTS & KEY FINDINGS', level=1)
heading4.runs[0].font.size = Pt(14)

para6 = doc.add_paragraph()
para6.style.font.size = Pt(11)
para6.add_run('The backtesting system successfully generated comprehensive performance analytics:\n\n')

results = [
    'Backtest Period: June 30, 2016 - August 31, 2025 (9.2 years)',
    'Total Return: 76.5% (using adjusted close prices)',
    'Compound Annual Growth Rate (CAGR): 6.43%',
    'Sharpe Ratio: 0.36 (indicating moderate risk-adjusted returns)',
    'Maximum Drawdown: -33.98% (occurred during COVID-19 crash in March 2020)',
    'Quarterly Hit Ratio: 38.5% (15 out of 39 quarters positive)',
    'Total Trades Executed: 1,800+ transactions over 39 rebalancing periods',
    'Average Momentum Score: Top selected stocks showed 15-25% returns over lookback period'
]

for result in results:
    para6.add_run(f'• {result}\n')

# 4.1 Key Insights
subheading4 = doc.add_heading('4.1 Key Insights', level=2)
subheading4.runs[0].font.size = Pt(12)

para7 = doc.add_paragraph()
para7.style.font.size = Pt(11)

insights = [
    'Price Type Impact: Using adjusted close prices yielded 16% higher returns than raw close prices (76.5% vs. 60.4%), highlighting the importance of corporate action adjustments',
    'Volatility Events: The strategy experienced significant drawdown during the 2020 pandemic but recovered over subsequent quarters',
    'Momentum Persistence: Selected stocks demonstrated varying performance, with ticker-level hit ratios ranging from 35-45%',
    'Quarterly Patterns: Portfolio value showed clear quarterly inflection points corresponding to rebalancing dates',
    'Transaction Costs: Total fees and slippage reduced returns by approximately 150 basis points annually'
]

for insight in insights:
    para7.add_run(f'• {insight}\n')

# 5. TECHNICAL DELIVERABLES
heading5 = doc.add_heading('5. TECHNICAL DELIVERABLES', level=1)
heading5.runs[0].font.size = Pt(14)

para8 = doc.add_paragraph()
para8.style.font.size = Pt(11)

deliverables = [
    'Core Backtesting Engine: 9 Python modules with 2,000+ lines of production code',
    'Test Suite: 70+ unit tests ensuring code reliability and accuracy',
    'Interactive Web Application: Streamlit-based interface with real-time visualizations',
    'Performance Visualizations: Equity curve, drawdown chart, rolling Sharpe ratio, quarterly heatmap',
    'Data Export Capabilities: CSV downloads for trades, holdings, returns, and quarterly selections',
    'Comprehensive Documentation: Code documentation, README, and deployment guides',
    'Version Control: Git repository with 25+ commits tracking development progress'
]

for deliverable in deliverables:
    para8.add_run(f'• {deliverable}\n')

# 6. SKILLS & TECHNOLOGIES UTILIZED
heading6 = doc.add_heading('6. SKILLS & TECHNOLOGIES UTILIZED', level=1)
heading6.runs[0].font.size = Pt(14)

para9 = doc.add_paragraph()
para9.style.font.size = Pt(11)

para9.add_run('Programming & Libraries: ').bold = True
para9.add_run('Python 3.13, Pandas, NumPy, Plotly, Streamlit\n')

para9.add_run('Financial Concepts: ').bold = True
para9.add_run('Momentum investing, backtesting methodology, risk metrics, portfolio management\n')

para9.add_run('Software Engineering: ').bold = True
para9.add_run('Object-oriented design, modular architecture, unit testing, version control (Git/GitHub)\n')

para9.add_run('Data Processing: ').bold = True
para9.add_run('Time series analysis, data validation, matrix operations, date parsing\n')

para9.add_run('Web Development: ').bold = True
para9.add_run('Interactive UI design, caching strategies, session state management\n')

# 7. CONCLUSION & LEARNING OUTCOMES
heading7 = doc.add_heading('7. CONCLUSION & LEARNING OUTCOMES', level=1)
heading7.runs[0].font.size = Pt(14)

para10 = doc.add_paragraph(
    'This internship project provided invaluable hands-on experience in quantitative finance and software development. '
    'The successful implementation of a production-grade backtesting system demonstrated the practical application of '
    'financial theory, data science, and software engineering principles. Key learnings include the critical importance '
    'of accurate data handling, the nuances of portfolio accounting, and the value of comprehensive testing in financial '
    'applications.\n\n'
    'The project reinforced the understanding that seemingly small bugs (such as the cash accounting error) can have '
    'dramatic impacts on calculated returns, emphasizing the need for rigorous validation. Additionally, the experience '
    'of building an end-to-end system—from data ingestion to user interface—provided insights into the full software '
    'development lifecycle in a financial context.\n\n'
    'Moving forward, the system can be extended with additional strategy types, risk management features, and real-time '
    'data integration, providing a solid foundation for future quantitative research and analysis.'
)
para10.style.font.size = Pt(11)

doc.add_paragraph()

# Signature section
sig_para = doc.add_paragraph()
sig_para.add_run('\n\n')
sig_para.add_run('_' * 30 + '\n')
sig_para.add_run('Student Signature\n')
sig_para.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}')
sig_para.style.font.size = Pt(11)

# Save document
doc.save('Internship_Report_Momentum_Backtester.docx')
print("✓ Internship report created successfully: Internship_Report_Momentum_Backtester.docx")
